import pandas as pd
from tkinter import Tk, Button, Label, filedialog, messagebox, StringVar
from tkinter.ttk import Progressbar
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import time

# 匯入文件的全局變量
file_path = ""


# 匯入文件的函數
def import_file():
    global file_path
    progress["value"] = 0
    progress_label_var.set("")
    file_path = filedialog.askopenfilename(
        title="選擇 Excel 文件",
        filetypes=[("Excel Files", "*.xlsx")]
    )
    if file_path:
        messagebox.showinfo("成功", "已成功匯入 Excel 文件")


# 輸出文件的函數
def export_to_word():
    if not file_path:
        messagebox.showwarning("警告", "請先匯入 Excel 文件")
        return

    progress["value"] = 0
    progress_label_var.set("")
    try:
        # 讀取 Excel 文件
        excel_data = pd.read_excel(file_path)

        # 獲取資料行數，來設置進度條的最大值
        total_rows = len(excel_data)

        # 設置進度條
        progress["maximum"] = total_rows
        progress_label_var.set("處理中...")

        # 創建 Word 文件
        doc = Document()

        # 添加寄件人資訊
        sender_info = """寄件人
姓名：倍加能股份有限公司
地址：600嘉義市西區民生北路151號3樓
電話：05-224-2828"""
        #doc.add_paragraph(sender_info)

        # 添加每位收件人的資訊並更新進度條
        for index, row in excel_data.iterrows():
            recipient_info = f"""收件人
姓名：{row['收件人']}
地址：{row['地址']}
電話：0{row['手機']}"""
            paragraph1 = doc.add_paragraph(sender_info)
            for run in paragraph1.runs:
                run.font.size = Pt(16)
                run.bold = True
                run.font.name = 'Microsoft JhengHei'  # 設置字體
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')  # 設置中文字體
            paragraph2 = doc.add_paragraph(recipient_info)
            for run in paragraph2.runs:
                run.font.size = Pt(16)
                run.bold = True
                run.font.name = 'Microsoft JhengHei'  # 設置字體
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')  # 設置中文字體

            # 模擬處理時間以顯示進度條效果
            time.sleep(0.1)
            progress["value"] += 1
            progress.update()

        # 保存 Word 文件
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Files", "*.docx")],
            title="保存 Word 文件"
        )
        if save_path:
            doc.save(save_path)
            messagebox.showinfo("成功", "已成功輸出至 Word 文件")
            progress_label_var.set("完成！")
        else:
            progress_label_var.set("未保存")
    except Exception as e:
        messagebox.showerror("錯誤", f"處理文件時發生錯誤: {e}")
        progress_label_var.set("錯誤")


# 創建主窗口
root = Tk()
root.title("郵局寄件面單製作")
root.geometry("400x250")

# 標籤
label = Label(root, text="請匯入 Excel 文件並輸出為 Word 文件")
label.pack(pady=10)

# 匯入按鈕
import_button = Button(root, text="匯入 Excel 文件", command=import_file,width=15)
import_button.pack(pady=10)

# 輸出按鈕
export_button = Button(root, text="輸出至 Word 文件", command=export_to_word,width=15)
export_button.pack(pady=10)

# 進度條標籤
progress_label_var = StringVar()
progress_label = Label(root, textvariable=progress_label_var)
progress_label.pack(pady=10)

# 進度條
progress = Progressbar(root, orient="horizontal", length=300, mode="determinate")
progress.pack(pady=10)

# 啟動主循環
root.mainloop()
